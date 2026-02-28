"""Script timeline estimation and export helpers (txt/srt/docx)."""

from __future__ import annotations

from pathlib import Path
import re
import json
import zipfile

from docx import Document as DocxDocument

from app.models import DialogueLine


def _line_words(text: str) -> int:
    words = re.findall(r"\S+", text or "")
    return max(1, len(words))


def _chapter_title(text: str) -> str:
    words = re.findall(r"\S+", (text or "").strip())
    if not words:
        return "Глава"
    title = " ".join(words[:8])
    return title if len(words) <= 8 else f"{title}…"


def estimate_timeline(script: list[DialogueLine] | list[dict], audio_duration_sec: float | None = None) -> dict:
    """Estimate per-line timeline and chapters."""
    lines: list[DialogueLine] = [line if isinstance(line, DialogueLine) else DialogueLine(**line) for line in script]
    if not lines:
        return {"lines": [], "chapters": [], "total_duration_sec": 0.0}

    word_counts = [_line_words(line.text) for line in lines]
    if audio_duration_sec and audio_duration_sec > 0:
        total_words = sum(word_counts) or 1
        durations = [(wc / total_words) * audio_duration_sec for wc in word_counts]
    else:
        # fallback: ~2.8 words/sec with sane bounds
        durations = [max(1.4, min(16.0, wc / 2.8)) for wc in word_counts]
    total_duration = sum(durations)

    timeline_rows: list[dict] = []
    t = 0.0
    for i, (line, dur) in enumerate(zip(lines, durations)):
        start = t
        end = t + dur
        timeline_rows.append(
            {
                "index": i,
                "voice": line.voice,
                "text": line.text,
                "start_sec": round(start, 3),
                "end_sec": round(end, 3),
                "duration_sec": round(dur, 3),
            }
        )
        t = end

    chapters: list[dict] = []
    ch_start_idx = 0
    ch_start_t = 0.0
    for i, row in enumerate(timeline_rows):
        ch_len_lines = i - ch_start_idx + 1
        ch_len_time = row["end_sec"] - ch_start_t
        should_cut = ch_len_lines >= 6 or ch_len_time >= 90.0
        if should_cut:
            first_text = timeline_rows[ch_start_idx]["text"]
            chapters.append(
                {
                    "index": len(chapters) + 1,
                    "title": _chapter_title(first_text),
                    "line_from": ch_start_idx,
                    "line_to": i,
                    "start_sec": round(ch_start_t, 3),
                    "end_sec": row["end_sec"],
                }
            )
            ch_start_idx = i + 1
            ch_start_t = row["end_sec"]
    if ch_start_idx < len(timeline_rows):
        first_text = timeline_rows[ch_start_idx]["text"]
        chapters.append(
            {
                "index": len(chapters) + 1,
                "title": _chapter_title(first_text),
                "line_from": ch_start_idx,
                "line_to": len(timeline_rows) - 1,
                "start_sec": round(ch_start_t, 3),
                "end_sec": round(total_duration, 3),
            }
        )

    return {
        "lines": timeline_rows,
        "chapters": chapters,
        "total_duration_sec": round(total_duration, 3),
    }


def _format_srt_time(sec: float) -> str:
    millis = int(round(max(0.0, sec) * 1000))
    h = millis // 3_600_000
    millis %= 3_600_000
    m = millis // 60_000
    millis %= 60_000
    s = millis // 1000
    ms = millis % 1000
    return f"{h:02}:{m:02}:{s:02},{ms:03}"


def render_txt(script: list[DialogueLine] | list[dict]) -> str:
    lines: list[DialogueLine] = [line if isinstance(line, DialogueLine) else DialogueLine(**line) for line in script]
    return "\n".join(f"[{line.voice}] {line.text}" for line in lines)


def render_srt(script: list[DialogueLine] | list[dict], audio_duration_sec: float | None = None) -> str:
    timeline = estimate_timeline(script, audio_duration_sec=audio_duration_sec)
    blocks: list[str] = []
    for i, row in enumerate(timeline["lines"], start=1):
        blocks.append(
            "\n".join(
                [
                    str(i),
                    f"{_format_srt_time(row['start_sec'])} --> {_format_srt_time(row['end_sec'])}",
                    f"[{row['voice']}] {row['text']}",
                    "",
                ]
            )
        )
    return "\n".join(blocks).strip() + "\n"


def save_docx(script: list[DialogueLine] | list[dict], out_path: Path, title: str) -> Path:
    lines: list[DialogueLine] = [line if isinstance(line, DialogueLine) else DialogueLine(**line) for line in script]
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(f"[{line.voice}] ")
        run.bold = True
        p.add_run(line.text)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def save_report_docx(
    *,
    out_path: Path,
    title: str,
    summary: str | None,
    sources: list[dict] | None,
    script: list[DialogueLine] | list[dict] | None,
    metrics: dict | None,
) -> Path:
    """Save rich DOCX report: summary + sources + script + metrics."""
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    doc.add_heading("Саммари", level=2)
    doc.add_paragraph((summary or "Нет саммари").strip() or "Нет саммари")

    src = sources or []
    if src:
        doc.add_heading("Источники", level=2)
        for s in src[:20]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{s.get('chunk_id', 'chunk')}: ").bold = True
            p.add_run(str(s.get("text", "")))

    lines = script or []
    if lines:
        doc.add_heading("Скрипт", level=2)
        for line in lines:
            row = line if isinstance(line, DialogueLine) else DialogueLine(**line)
            p = doc.add_paragraph()
            p.add_run(f"[{row.voice}] ").bold = True
            p.add_run(row.text)

    if metrics:
        doc.add_heading("Метрики", level=2)
        doc.add_paragraph(json.dumps(metrics, ensure_ascii=False, indent=2))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def build_document_bundle(
    *,
    out_path: Path,
    document_id: str,
    filename: str,
    summary: str | None,
    sources: list[dict] | None,
    script: list[dict] | None,
    metrics: dict | None,
    timeline: dict | None,
    attachments: list[Path] | None = None,
) -> Path:
    """Create zip bundle with all available artifacts for one document."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    txt_script = render_txt(script or []) if script else ""
    srt_script = render_srt(script or []) if script else ""
    report_docx_name = f"{document_id}_report.docx"
    with zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        manifest = {
            "document_id": document_id,
            "filename": filename,
            "has_summary": bool(summary),
            "has_script": bool(script),
            "has_metrics": bool(metrics),
            "attachments": [p.name for p in (attachments or []) if p.exists()],
        }
        zf.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))
        if summary:
            zf.writestr("summary.md", summary)
        if sources:
            zf.writestr("sources.json", json.dumps(sources, ensure_ascii=False, indent=2))
        if script:
            zf.writestr("script.json", json.dumps({"script": script}, ensure_ascii=False, indent=2))
            zf.writestr("script.txt", txt_script)
            zf.writestr("script.srt", srt_script)
        if metrics:
            zf.writestr("script_metrics.json", json.dumps(metrics, ensure_ascii=False, indent=2))
        if timeline:
            zf.writestr("timeline.json", json.dumps(timeline, ensure_ascii=False, indent=2))

        # Temporary DOCX inside outputs and then add to zip
        if summary or script:
            temp_docx = out_path.parent / report_docx_name
            save_report_docx(
                out_path=temp_docx,
                title=f"Отчёт по документу {document_id}",
                summary=summary,
                sources=sources,
                script=script,
                metrics=metrics,
            )
            if temp_docx.exists():
                zf.write(temp_docx, arcname=report_docx_name)
                try:
                    temp_docx.unlink()
                except OSError:
                    pass

        for p in attachments or []:
            if p.exists() and p.is_file():
                zf.write(p, arcname=f"media/{p.name}")
    return out_path
